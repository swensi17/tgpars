from flask import Flask, render_template, request, redirect, url_for, flash, session, send_file, jsonify, make_response
from flask_login import LoginManager, UserMixin, login_user, login_required, logout_user, current_user
from telethon import TelegramClient, sync, functions, types
from telethon.tl.types import InputPeerChannel, InputPeerUser, InputPeerChat
from telethon.tl.functions.channels import JoinChannelRequest
from telethon.tl.functions.contacts import AddContactRequest
from telethon.errors import SessionPasswordNeededError, PhoneCodeInvalidError
import sqlite3
import os
import asyncio
import nest_asyncio
from dotenv import load_dotenv
from contextlib import contextmanager
from datetime import datetime
import random
from werkzeug.utils import secure_filename
import csv
import json
from docx.api import Document
from io import StringIO

# Применяем nest_asyncio для решения проблем с event loop
nest_asyncio.apply()

# Загрузка переменных окружения
load_dotenv()

app = Flask(__name__)
app.config['SECRET_KEY'] = os.getenv('FLASK_SECRET_KEY')
login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'login'
login_manager.login_message = 'Пожалуйста, войдите для доступа к этой странице.'

# Создаем папку для сессий, если она не существует
if not os.path.exists('sessions'):
    os.makedirs('sessions')

# Конфигурация Telegram API
API_ID = os.getenv('TELEGRAM_API_ID')
API_HASH = os.getenv('TELEGRAM_API_HASH')

def run_async(coro):
    loop = asyncio.new_event_loop()
    asyncio.set_event_loop(loop)
    try:
        return loop.run_until_complete(coro)
    finally:
        loop.close()

class TelegramConnection:
    def __init__(self, session_file):
        self.client = TelegramClient(session_file, API_ID, API_HASH)
        
    async def __aenter__(self):
        await self.client.connect()
        return self.client
        
    async def __aexit__(self, exc_type, exc_val, exc_tb):
        await self.client.disconnect()

async def with_client(session_file, operation):
    async with TelegramConnection(session_file) as client:
        return await operation(client)

class User(UserMixin):
    def __init__(self, id, phone):
        self.id = id
        self.phone = phone

@login_manager.user_loader
def load_user(user_id):
    conn = get_db_connection()
    user = conn.execute('SELECT * FROM users WHERE id = ?', (user_id,)).fetchone()
    conn.close()
    if user:
        return User(user['id'], user['phone'])
    return None

def get_db_connection():
    conn = sqlite3.connect('database.db')
    conn.row_factory = sqlite3.Row
    return conn

@app.route('/')
def index():
    if current_user.is_authenticated:
        return redirect(url_for('dashboard'))
    return redirect(url_for('login'))

@app.route('/login', methods=['GET', 'POST'])
def login():
    if current_user.is_authenticated:
        return redirect(url_for('dashboard'))
        
    if request.method == 'POST':
        phone = request.form['phone']
        session_file = f'sessions/{phone}.session'
        
        try:
            # Создаем пользователя в базе данных, если его еще нет
            conn = get_db_connection()
            user = conn.execute('SELECT * FROM users WHERE phone = ?', (phone,)).fetchone()
            if not user:
                conn.execute('INSERT INTO users (phone) VALUES (?)', (phone,))
                conn.commit()
            conn.close()

            async def login_operation(client):
                if not await client.is_user_authorized():
                    result = await client.send_code_request(phone)
                    session['phone'] = phone
                    session['phone_code_hash'] = result.phone_code_hash
                    return False
                return True

            is_authorized = run_async(with_client(session_file, login_operation))
            
            if not is_authorized:
                flash('Код подтверждения отправлен на ваш телефон', 'success')
                return redirect(url_for('verify_code'))
            else:
                conn = get_db_connection()
                user = conn.execute('SELECT * FROM users WHERE phone = ?', (phone,)).fetchone()
                if user:
                    user_obj = User(user['id'], user['phone'])
                    login_user(user_obj, remember=True)
                    conn.execute('UPDATE users SET is_authorized = 1 WHERE id = ?', (user['id'],))
                    conn.commit()
                conn.close()
                return redirect(url_for('dashboard'))
            
        except Exception as e:
            flash(f'Ошибка при подключении к Telegram: {str(e)}', 'error')
            return redirect(url_for('login'))
    
    return render_template('login.html')

@app.route('/verify_code', methods=['GET', 'POST'])
def verify_code():
    if current_user.is_authenticated:
        return redirect(url_for('dashboard'))
        
    if 'phone' not in session or 'phone_code_hash' not in session:
        return redirect(url_for('login'))
        
    if request.method == 'POST':
        code = request.form['code']
        phone = session['phone']
        phone_code_hash = session['phone_code_hash']
        session_file = f'sessions/{phone}.session'
        
        try:
            async def verify_operation(client):
                try:
                    await client.sign_in(phone=phone, code=code, phone_code_hash=phone_code_hash)
                    return True, None
                except SessionPasswordNeededError:
                    return False, 'two_factor'
                except PhoneCodeInvalidError:
                    return False, 'invalid_code'
            
            success, error = run_async(with_client(session_file, verify_operation))
            
            if success:
                conn = get_db_connection()
                user = conn.execute('SELECT * FROM users WHERE phone = ?', (phone,)).fetchone()
                if user:
                    user_obj = User(user['id'], user['phone'])
                    login_user(user_obj, remember=True)
                    conn.execute('UPDATE users SET is_authorized = 1 WHERE id = ?', (user['id'],))
                    conn.commit()
                conn.close()
                
                session.pop('phone', None)
                session.pop('phone_code_hash', None)
                flash('Успешная авторизация в Telegram', 'success')
                return redirect(url_for('dashboard'))
            elif error == 'two_factor':
                session['two_factor_auth'] = True
                flash('Требуется двухфакторная аутентификация', 'info')
                return redirect(url_for('two_factor_auth'))
            else:
                flash('Неверный код подтверждения', 'error')
                
        except Exception as e:
            flash(f'Ошибка: {str(e)}', 'error')
            
    return render_template('verify_code.html')

@app.route('/two_factor_auth', methods=['GET', 'POST'])
def two_factor_auth():
    if current_user.is_authenticated:
        return redirect(url_for('dashboard'))
        
    if 'two_factor_auth' not in session or 'phone' not in session:
        return redirect(url_for('login'))
        
    if request.method == 'POST':
        password = request.form['password']
        phone = session['phone']
        session_file = f'sessions/{phone}.session'
        
        try:
            async def two_factor_operation(client):
                await client.sign_in(password=password)
                return True
            
            success = run_async(with_client(session_file, two_factor_operation))
            
            if success:
                conn = get_db_connection()
                user = conn.execute('SELECT * FROM users WHERE phone = ?', (phone,)).fetchone()
                if user:
                    user_obj = User(user['id'], user['phone'])
                    login_user(user_obj, remember=True)
                    conn.execute('UPDATE users SET is_authorized = 1 WHERE id = ?', (user['id'],))
                    conn.commit()
                conn.close()
                
                session.pop('phone', None)
                session.pop('two_factor_auth', None)
                
                flash('Успешная авторизация в Telegram', 'success')
                return redirect(url_for('dashboard'))
                
        except Exception as e:
            flash(f'Ошибка: {str(e)}', 'error')
    
    return render_template('two_factor_auth.html')

@app.route('/logout')
@login_required
def logout():
    if current_user.is_authenticated:
        conn = get_db_connection()
        conn.execute('UPDATE users SET is_authorized = 0 WHERE id = ?', (current_user.id,))
        conn.commit()
        conn.close()
    logout_user()
    flash('Вы успешно вышли из системы', 'success')
    return redirect(url_for('login'))

@app.route('/dashboard')
@login_required
def dashboard():
    return render_template('dashboard.html')

@app.route('/parse_channel', methods=['GET', 'POST'])
@login_required
def parse_channel():
    if request.method == 'POST':
        channel = request.form['channel']
        limit = int(request.form.get('limit', 100))
        offset = int(request.form.get('offset', 0))
        batch_size = int(request.form.get('batch_size', 50))
        date_from = request.form.get('date_from')
        date_to = request.form.get('date_to')
        keywords = request.form.get('keywords', '').split(',') if request.form.get('keywords') else []
        include_media = 'include_media' in request.form
        include_forwarded = 'include_forwarded' in request.form
        include_links = 'include_links' in request.form
        include_views = 'include_views' in request.form
        
        try:
            session_file = f'sessions/{current_user.phone}.session'
            
            if not os.path.exists(session_file):
                flash('Необходимо сначала авторизоваться в Telegram', 'error')
                return redirect(url_for('login'))
                
            # Сохраняем параметры парсинга в сессии
            session['parsing_params'] = {
                'total_count': limit,
                'processed_count': 0,
                'filtered_count': 0,
                'is_running': True
            }
            
            async def parse_operation(client):
                if not await client.is_user_authorized():
                    flash('Необходима авторизация в Telegram', 'error')
                    return None
                    
                messages = []
                processed_count = 0
                filtered_count = 0
                
                # Проверяем корректность канала
                try:
                    entity = await client.get_entity(channel)
                    if not hasattr(entity, 'title'):
                        flash('Указанный канал не найден или недоступен', 'error')
                        return None
                except ValueError:
                    flash('Неверный формат канала. Используйте @username или полную ссылку', 'error')
                    return None
                except Exception as e:
                    flash(f'Ошибка при доступе к каналу: {str(e)}', 'error')
                    return None
                
                # Проверяем параметры
                if limit <= 0 or limit > 5000:
                    flash('Количество сообщений должно быть от 1 до 5000', 'error')
                    return None
                
                if offset < 0:
                    flash('Смещение не может быть отрицательным', 'error')
                    return None
                
                # Получаем общее количество сообщений для прогресса
                try:
                    async for _ in client.iter_messages(channel, limit=1):
                        total_messages = _.id
                        break
                except Exception as e:
                    flash(f'Ошибка при получении информации о канале: {str(e)}', 'error')
                    return None
                
                # Обновляем общее количество сообщений
                session['parsing_params']['total_count'] = min(limit, total_messages - offset)
                
                try:
                    async for message in client.iter_messages(
                        channel, 
                        limit=limit,
                        offset_id=offset,
                        reverse=True
                    ):
                        # Проверяем, не остановлен ли парсинг
                        if not session.get('parsing_params', {}).get('is_running', True):
                            flash('Парсинг был остановлен пользователем', 'info')
                            break
                            
                        processed_count += 1
                        session['parsing_params']['processed_count'] = processed_count
                        
                        # Обновляем прогресс
                        progress = int((processed_count / limit) * 100)
                        session['parse_progress'] = progress
                        
                        try:
                            # Проверяем дату
                            if date_from:
                                try:
                                    date_from_obj = datetime.strptime(date_from, '%Y-%m-%d').date()
                                    if message.date.date() < date_from_obj:
                                        filtered_count += 1
                                        continue
                                except ValueError:
                                    flash('Неверный формат даты начала', 'error')
                                    return None
                                    
                            if date_to:
                                try:
                                    date_to_obj = datetime.strptime(date_to, '%Y-%m-%d').date()
                                    if message.date.date() > date_to_obj:
                                        filtered_count += 1
                                        continue
                                except ValueError:
                                    flash('Неверный формат даты окончания', 'error')
                                    return None
                            
                            # Проверяем пересланные сообщения
                            if not include_forwarded and message.forward:
                                filtered_count += 1
                                continue
                            
                            # Проверяем медиафайлы
                            if not include_media and message.media:
                                filtered_count += 1
                                continue
                            
                            # Проверяем ссылки
                            has_links = False
                            if message.entities:
                                for entity in message.entities:
                                    if hasattr(entity, 'url') or str(type(entity).__name__) in ['MessageEntityUrl', 'MessageEntityTextUrl']:
                                        has_links = True
                                        break
                                        
                            if not include_links and has_links:
                                filtered_count += 1
                                continue
                            
                            # Проверяем ключевые слова
                            if keywords:
                                message_text = message.text or ''
                                if not any(keyword.strip().lower() in message_text.lower() for keyword in keywords if keyword.strip()):
                                    filtered_count += 1
                                    continue
                            
                            # Обновляем счетчик отфильтрованных
                            session['parsing_params']['filtered_count'] = filtered_count
                            
                            # Добавляем сообщение в результаты
                            messages.append({
                                'id': message.id,
                                'date': message.date.strftime('%Y-%m-%d %H:%M:%S'),
                                'text': message.text or '',
                                'has_media': bool(message.media),
                                'is_forwarded': bool(message.forward),
                                'has_links': has_links,
                                'views': message.views if include_views and hasattr(message, 'views') else None
                            })
                            
                        except Exception as e:
                            flash(f'Ошибка при обработке сообщения {message.id}: {str(e)}', 'error')
                            continue
                            
                except Exception as e:
                    flash(f'Ошибка при парсинге сообщений: {str(e)}', 'error')
                    return None
                
                if not messages:
                    flash('Не найдено сообщений, соответствующих заданным критериям', 'warning')
                
                return messages
                
            messages = run_async(with_client(session_file, parse_operation))
            
            # Очищаем параметры парсинга
            session.pop('parsing_params', None)
            
            if messages is None:
                flash('Ошибка при парсинге канала', 'error')
                return redirect(url_for('parse_channel'))
                
            return render_template('parse_results.html', messages=messages, channel=channel)
            
        except Exception as e:
            flash(f'Ошибка при парсинге канала: {str(e)}', 'error')
            
    return render_template('parse_channel.html')

@app.route('/send_message', methods=['GET', 'POST'])
@login_required
def send_message():
    if request.method == 'POST':
        # Получаем основные параметры
        recipients = request.form['recipients'].split('\n')
        message = request.form['message']
        delay = int(request.form.get('delay', 2))
        batch_size = int(request.form.get('batch_size', 10))
        repeat_count = int(request.form.get('repeat_count', 1))
        interval = int(request.form.get('interval', 0))
        start_time = int(request.form.get('start_time', 0))
        
        # Получаем дополнительные параметры
        disable_preview = 'disable_preview' in request.form
        silent = 'silent' in request.form
        random_delay = 'random_delay' in request.form
        shuffle_recipients = 'shuffle_recipients' in request.form
        
        # Получаем инлайн кнопки
        button_texts = request.form.getlist('button_text[]')
        button_urls = request.form.getlist('button_url[]')
        inline_buttons = []
        for text, url in zip(button_texts, button_urls):
            if text.strip() and url.strip():
                inline_buttons.append({'text': text.strip(), 'url': url.strip()})
        
        # Получаем файл, если он был загружен
        media_file = request.files.get('media_file')
        caption = request.form.get('caption', '')
        
        # Очищаем список получателей
        recipients = [r.strip() for r in recipients if r.strip()]
        
        if not recipients:
            flash('Список получателей пуст', 'error')
            return render_template('send_message.html', form_data=request.form)
            
        if not message and not media_file and not caption:
            flash('Необходимо указать текст сообщения или загрузить файл', 'error')
            return render_template('send_message.html', form_data=request.form)
            
        try:
            session_file = f'sessions/{current_user.phone}.session'
            
            if not os.path.exists(session_file):
                flash('Необходимо сначала авторизоваться в Telegram', 'error')
                return redirect(url_for('login'))
                
            # Сохраняем файл, если он был загружен
            media_path = None
            if media_file:
                filename = secure_filename(media_file.filename)
                media_path = os.path.join('uploads', filename)
                os.makedirs('uploads', exist_ok=True)
                media_file.save(media_path)
                
            # Сохраняем параметры отправки в сессии
            total_messages = len(recipients) * repeat_count
            session['sending_params'] = {
                'total_count': total_messages,
                'success_count': 0,
                'error_count': 0,
                'remaining_count': total_messages,
                'is_running': True,
                'last_status': 'Подготовка к отправке...'
            }
            
            async def send_operation(client):
                if not await client.is_user_authorized():
                    flash('Необходима авторизация в Telegram', 'error')
                    return False, 0, 0
                    
                success_count = 0
                error_count = 0
                total = total_messages
                remaining_count = total
                
                # Создаем инлайн кнопки для Telegram
                buttons = None
                if inline_buttons:
                    buttons = []
                    for btn in inline_buttons:
                        buttons.append([types.KeyboardButtonUrl(btn['text'], btn['url'])])
                
                try:
                    # Определяем режим отправки
                    sending_mode = request.form.get('sending_mode', 'parallel')
                    
                    # Основной цикл отправки
                    for repeat in range(repeat_count):
                        current_recipients = recipients.copy()
                        if shuffle_recipients:
                            random.shuffle(current_recipients)
                        
                        for recipient in current_recipients:
                            if not session.get('sending_params', {}).get('is_running', True):
                                session['sending_params'].update({
                                    'last_status': 'Рассылка остановлена пользователем',
                                    'is_running': False,
                                    'remaining_count': total - (success_count + error_count)
                                })
                                return False, success_count, error_count
                            
                            try:
                                session['sending_params'].update({
                                    'current_recipient': recipient,
                                    'last_status': f'Отправка для {recipient}...'
                                })
                                
                                if media_path:
                                    await client.send_file(
                                        recipient,
                                        media_path,
                                        caption=caption or message,
                                        parse_mode='md',
                                        buttons=buttons if buttons else None,
                                        silent=silent
                                    )
                                else:
                                    await client.send_message(
                                        recipient,
                                        message,
                                        link_preview=not disable_preview,
                                        buttons=buttons if buttons else None,
                                        silent=silent,
                                        parse_mode='md'
                                    )
                                
                                success_count += 1
                                remaining_count -= 1
                                
                            except Exception as e:
                                error_count += 1
                                remaining_count -= 1
                                error_message = str(e)
                                if 'Too many requests' in error_message:
                                    error_message = 'Слишком много запросов, нужно подождать'
                                elif 'privacy' in error_message.lower():
                                    error_message = 'Пользователь ограничил отправку сообщений'
                                elif 'not found' in error_message.lower():
                                    error_message = 'Пользователь не найден'
                                
                                session['sending_params'].update({
                                    'error_count': error_count,
                                    'last_status': f'Ошибка отправки для {recipient}: {error_message}'
                                })
                            
                            # Обновляем статус
                            session['sending_params'].update({
                                'success_count': success_count,
                                'send_progress': int((success_count + error_count) / total * 100),
                                'remaining_count': remaining_count
                            })
                            
                            # Задержка между сообщениями
                            if delay > 0:
                                actual_delay = random.uniform(delay * 0.5, delay * 1.5) if random_delay else delay
                                session['sending_params']['last_status'] = f'Ожидание {actual_delay:.1f} сек перед следующей отправкой...'
                                await asyncio.sleep(actual_delay)
                        
                        # Интервал между повторами
                        if interval > 0 and repeat < repeat_count - 1:
                            session['sending_params']['last_status'] = f'Ожидание {interval} минут перед следующим повтором...'
                            await asyncio.sleep(interval * 60)
                    
                    # Удаляем временный файл
                    if media_path and os.path.exists(media_path):
                        os.remove(media_path)
                    
                    final_status = 'Рассылка завершена. '
                    if success_count > 0:
                        final_status += f'Успешно отправлено: {success_count}. '
                    if error_count > 0:
                        final_status += f'Ошибок: {error_count}. '
                    
                    session['sending_params'].update({
                        'last_status': final_status.strip(),
                        'is_running': False,
                        'send_progress': 100,
                        'remaining_count': 0
                    })
                    return True, success_count, error_count
                    
                except Exception as e:
                    session['sending_params'].update({
                        'last_status': f'Критическая ошибка: {str(e)}',
                        'is_running': False,
                        'remaining_count': total - (success_count + error_count)
                    })
                    return False, success_count, error_count
            
            success, success_count, error_count = run_async(with_client(session_file, send_operation))
            
            return render_template('send_message.html', 
                                form_data=request.form,
                                show_status=True,
                                success_count=success_count,
                                error_count=error_count)
            
        except Exception as e:
            flash(f'Ошибка при отправке сообщений: {str(e)}', 'error')
            return render_template('send_message.html', form_data=request.form)
            
    return render_template('send_message.html')

@app.route('/progress')
def progress():
    """Endpoint для получения прогресса операций"""
    sending_params = session.get('sending_params', {})
    parsing_params = session.get('parsing_params', {})
    return {
        'parse_progress': session.get('parse_progress', 0),
        'send_progress': session.get('send_progress', 0),
        'success_count': sending_params.get('success_count', 0),
        'error_count': sending_params.get('error_count', 0),
        'total_count': sending_params.get('total_count', 0),
        'remaining_count': sending_params.get('remaining_count', 0),
        'last_status': sending_params.get('last_status', ''),
        'current_channel': sending_params.get('current_channel', ''),
        'current_recipient': sending_params.get('current_recipient', ''),
        'is_running': sending_params.get('is_running', False),
        'processed_count': parsing_params.get('processed_count', 0),
        'filtered_count': parsing_params.get('filtered_count', 0),
        'parsing_total': parsing_params.get('total_count', 0)
    }

@app.route('/stop_parsing', methods=['POST'])
@login_required
def stop_parsing():
    """Endpoint для остановки парсинга"""
    if 'parsing_params' in session:
        session['parsing_params']['is_running'] = False
    return {'success': True}

@app.route('/stop_sending', methods=['POST'])
@login_required
def stop_sending():
    """Endpoint для остановки отправки"""
    if 'sending_params' in session:
        session['sending_params']['is_running'] = False
    return {'success': True}

@app.route('/global_search', methods=['GET', 'POST'])
@login_required
def global_search():
    if request.method == 'POST':
        query = request.form.get('query', '')
        search_terms = [term.strip() for term in query.split(',') if term.strip()]
        search_type = request.form.getlist('search_type')  # channels, chats, bots
        
        # Get filter parameters
        search_in_username = 'search_in_username' in request.form
        search_in_title = 'search_in_title' in request.form
        search_in_description = 'search_in_description' in request.form
        min_members = request.form.get('min_members', '')
        max_members = request.form.get('max_members', '')
        
        # Convert member counts to integers if provided
        try:
            min_members = int(min_members) if min_members else None
            max_members = int(max_members) if max_members else None
        except ValueError:
            min_members = None
            max_members = None
        
        try:
            session_file = f'sessions/{current_user.phone}.session'
            
            if not os.path.exists(session_file):
                flash('Необходимо сначала авторизоваться в Telegram', 'error')
                return redirect(url_for('login'))
            
            async def search_operation(client):
                results = {
                    'channels': [],
                    'chats': [],
                    'bots': []
                }
                
                try:
                    # Get list of subscribed channels and chats
                    dialogs = await client.get_dialogs()
                    subscribed_entities = {
                        dialog.entity.username.lower(): dialog.entity.id
                        for dialog in dialogs
                        if hasattr(dialog.entity, 'username') and dialog.entity.username
                    }
                    
                    # Get list of bot contacts
                    contacts = await client(functions.contacts.GetContactsRequest(hash=0))
                    bot_contacts = {
                        user.username.lower(): user.id
                        for user in contacts.users
                        if user.bot and user.username
                    }

                    all_results = {}
                    for term in search_terms:
                        search_results = await client(functions.contacts.SearchRequest(
                            q=term,
                            limit=100
                        ))
                        
                        # Process channels and chats
                        for chat in search_results.chats:
                            chat_id = str(chat.id)
                            if chat_id not in all_results:
                                matches = []
                                
                                # Check if entity matches search criteria
                                if search_in_username and hasattr(chat, 'username') and chat.username:
                                    if any(t.lower() in chat.username.lower() for t in search_terms):
                                        matches.append('username')
                                        
                                if search_in_title and hasattr(chat, 'title'):
                                    if any(t.lower() in chat.title.lower() for t in search_terms):
                                        matches.append('title')
                                        
                                if search_in_description and hasattr(chat, 'about') and chat.about:
                                    if any(t.lower() in chat.about.lower() for t in search_terms):
                                        matches.append('description')
                                
                                # Check members count
                                members_ok = True
                                if hasattr(chat, 'participants_count'):
                                    if min_members and chat.participants_count < min_members:
                                        members_ok = False
                                    if max_members and chat.participants_count > max_members:
                                        members_ok = False

                                # Check subscription status
                                is_subscribed = (
                                    hasattr(chat, 'username') and 
                                    chat.username and 
                                    chat.username.lower() in subscribed_entities
                                )
                                
                                # Only add to results if matches found and members count is ok
                                if (matches or not any([search_in_username, search_in_title, search_in_description])) and members_ok:
                                    if hasattr(chat, 'broadcast') and chat.broadcast:
                                        all_results[chat_id] = {
                                            'type': 'channel',
                                            'data': {
                                                'id': chat.id,
                                                'title': chat.title,
                                                'username': chat.username if hasattr(chat, 'username') else None,
                                                'members': chat.participants_count if hasattr(chat, 'participants_count') else None,
                                                'description': chat.about if hasattr(chat, 'about') else None,
                                                'matched_terms': matches,
                                                'is_subscribed': is_subscribed
                                            }
                                        }
                                    elif hasattr(chat, 'megagroup') and chat.megagroup:
                                        all_results[chat_id] = {
                                            'type': 'chat',
                                            'data': {
                                                'id': chat.id,
                                                'title': chat.title,
                                                'username': chat.username if hasattr(chat, 'username') else None,
                                                'members': chat.participants_count if hasattr(chat, 'participants_count') else None,
                                                'description': chat.about if hasattr(chat, 'about') else None,
                                                'matched_terms': matches,
                                                'is_subscribed': is_subscribed
                                            }
                                        }

                        # Process bots
                        for user in search_results.users:
                            if user.bot:
                                user_id = str(user.id)
                                if user_id not in all_results:
                                    matches = []
                                    if search_in_username and user.username:
                                        if any(t.lower() in user.username.lower() for t in search_terms):
                                            matches.append('username')
                                    if search_in_title and user.first_name:
                                        if any(t.lower() in user.first_name.lower() for t in search_terms):
                                            matches.append('name')
                                    if search_in_description and hasattr(user, 'about') and user.about:
                                        if any(t.lower() in user.about.lower() for t in search_terms):
                                            matches.append('description')
                                    
                                    # Check if bot is in contacts
                                    is_subscribed = user.username and user.username.lower() in bot_contacts
                                    
                                    if matches or not any([search_in_username, search_in_title, search_in_description]):
                                        all_results[user_id] = {
                                            'type': 'bot',
                                            'data': {
                                                'id': user.id,
                                                'username': user.username,
                                                'name': f"{user.first_name} {user.last_name if user.last_name else ''}".strip(),
                                                'description': user.about if hasattr(user, 'about') else None,
                                                'matched_terms': matches,
                                                'is_subscribed': is_subscribed
                                            }
                                        }

                    # Organize results by type
                    for item_id, item_data in all_results.items():
                        if item_data['type'] == 'channel' and ('channels' in search_type or not search_type):
                            results['channels'].append(item_data['data'])
                        elif item_data['type'] == 'chat' and ('chats' in search_type or not search_type):
                            results['chats'].append(item_data['data'])
                        elif item_data['type'] == 'bot' and ('bots' in search_type or not search_type):
                            results['bots'].append(item_data['data'])

                    # Sort results by number of matched terms
                    for category in results:
                        results[category].sort(key=lambda x: len(x.get('matched_terms', [])), reverse=True)
                    
                    return results
                    
                except Exception as e:
                    flash(f'Ошибка при поиске: {str(e)}', 'error')
                    return None
            
            results = run_async(with_client(session_file, search_operation))
            if results:
                return render_template('global_search.html', 
                                    results=results, 
                                    query=query,
                                    search_type=search_type,
                                    search_in_username=search_in_username,
                                    search_in_title=search_in_title,
                                    search_in_description=search_in_description,
                                    min_members=min_members,
                                    max_members=max_members)
            
        except Exception as e:
            flash(f'Ошибка при поиске: {str(e)}', 'error')
            
    return render_template('global_search.html')

@app.route('/unsubscribe/<entity_type>/<username>', methods=['POST'])
@login_required
def unsubscribe(entity_type, username):
    try:
        session_file = f'sessions/{current_user.phone}.session'
        
        async def unsubscribe_operation(client):
            try:
                entity = await client.get_entity(f"@{username}")
                await client(functions.channels.LeaveChannelRequest(entity))
                return True
            except Exception as e:
                error_msg = str(e)
                if "CHAT_NOT_MODIFIED" in error_msg:
                    return "Вы уже отписаны"
                return error_msg
                
        result = run_async(with_client(session_file, unsubscribe_operation))
        return jsonify({'success': result == True, 'error': result if result != True else None})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@app.route('/subscribe/<entity_type>/<username>', methods=['POST'])
@login_required
def subscribe(entity_type, username):
    try:
        session_file = f'sessions/{current_user.phone}.session'
        
        async def subscribe_operation(client):
            try:
                # Get the entity first
                try:
                    entity = await client.get_entity(username)
                except ValueError:
                    try:
                        entity = await client.get_entity(f"@{username}")
                    except ValueError:
                        return "Канал или чат не найден"
                
                # Join using the appropriate method based on entity type
                if entity_type in ['channels', 'chats']:
                    try:
                        await client(JoinChannelRequest(entity))
                        return True
                    except Exception as e:
                        if "CHANNELS_TOO_MUCH" in str(e):
                            return "Достигнут лимит каналов"
                        elif "CHAT_WRITE_FORBIDDEN" in str(e):
                            return "Доступ запрещен"
                        elif "FLOOD_WAIT" in str(e):
                            return "Слишком много запросов, попробуйте позже"
                        else:
                            return str(e)
                elif entity_type == 'bots':
                    try:
                        await client(AddContactRequest(
                            id=entity,
                            first_name="Bot",
                            last_name="",
                            phone="",
                            add_phone_privacy_exception=False
                        ))
                        return True
                    except Exception as e:
                        return str(e)
                return True
            except Exception as e:
                error_msg = str(e)
                if "FLOOD_WAIT" in error_msg:
                    return "Слишком много запросов, попробуйте позже"
                return error_msg
                
        result = run_async(with_client(session_file, subscribe_operation))
        return jsonify({'success': result == True, 'error': result if result != True else None})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@app.route('/export_results/<format>', methods=['POST'])
@login_required
def export_results(format):
    try:
        data = request.json
        if format == 'csv':
            output = StringIO()
            writer = csv.writer(output, dialect='excel')  # Remove encoding parameter as it's not supported
            writer.writerow(['Тип', 'Название', 'Юзернейм', 'Описание', 'Участники'])
            
            for type_key in ['channels', 'chats', 'bots']:
                for item in data.get(type_key, []):
                    row = [
                        type_key,
                        item.get('title', item.get('name', '')),
                        item.get('username', ''),
                        item.get('description', ''),
                        item.get('members', '')
                    ]
                    writer.writerow([str(cell).encode('utf-8').decode('utf-8') for cell in row])
            
            output.seek(0)
            response = make_response(output.getvalue().encode('utf-8-sig'))
            response.headers['Content-Type'] = 'text/csv; charset=utf-8-sig'
            response.headers['Content-Disposition'] = 'attachment; filename=search_results.csv'
            return response
            
        elif format == 'txt':
            output = StringIO()
            for type_key in ['channels', 'chats', 'bots']:
                output.write(f"\n=== {type_key.upper()} ===\n\n")
                for item in data.get(type_key, []):
                    output.write(f"Название: {item.get('title', item.get('name', ''))}\n")
                    output.write(f"Юзернейм: @{item.get('username', '')}\n")
                    if item.get('description'):
                        output.write(f"Описание: {item.get('description')}\n")
                    if item.get('members'):
                        output.write(f"Участники: {item.get('members')}\n")
                    output.write("\n")
            
            output.seek(0)
            response = make_response(output.getvalue())
            response.headers['Content-Type'] = 'text/plain; charset=utf-8'
            response.headers['Content-Disposition'] = 'attachment; filename=search_results.txt'
            return response

            
        elif format == 'docx':
            doc = Document()
            for type_key in ['channels', 'chats', 'bots']:
                doc.add_heading(f"{type_key.upper()}", level=1)
                for item in data.get(type_key, []):
                    doc.add_heading(item.get('title', item.get('name', '')), level=2)
                    doc.add_paragraph(f"Юзернейм: @{item.get('username', '')}")
                    if item.get('description'):
                        doc.add_paragraph(f"Описание: {item.get('description')}")
                    if item.get('members'):
                        doc.add_paragraph(f"Участники: {item.get('members')}")
                    doc.add_paragraph("\n")
            
            temp_path = 'temp_results.docx'
            doc.save(temp_path)
            return send_file(
                temp_path,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                as_attachment=True,
                download_name='search_results.docx'
            )
            
        return jsonify({'error': 'Неподдерживаемый формат'}), 400
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    app.run(debug=True) 